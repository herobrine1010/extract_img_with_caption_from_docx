import os
from docx import Document
from PIL import Image
from io import BytesIO
from lxml import etree


def contains_image(run):
    """检查 run 是否包含图片"""
    # 获取当前 run 的 XML 字符串
    run_element_str = run._element.xml
    run_element = etree.fromstring(run_element_str)

    # 定义命名空间
    namespaces = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',  # drawingML 命名空间（用于图片）
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',  # 关系命名空间
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'  # Word 处理命名空间
    }

    # 查找 <a:blip> 标签（通常表示内嵌图片）
    inline_shapes = run_element.xpath('.//a:blip', namespaces=namespaces)

    # 查找 <w:object> 标签（可能表示嵌入对象，如图片、图形等）
    object_shapes = run_element.xpath('.//w:object', namespaces=namespaces)

    # 如果找到了 <a:blip> 或 <w:object> 标签，返回 True，表示包含图片
    return bool(inline_shapes or object_shapes)


def convert_emf_or_wmf(image_data, output_path):
    """尝试将 EMF/WMF 图片转换为 PNG 格式并保存。"""
    try:
        with Image.open(BytesIO(image_data)) as img:
            img.save(output_path, "PNG")
            print(f"图片已保存并转换为 PNG 格式：{output_path}")
    except Exception as e:
        print(f"无法处理图片格式，跳过该图片: {e}")


def extract_images_and_captions(docx_path, output_dir):
    os.makedirs(output_dir, exist_ok=True)
    formula_dir = os.path.join(output_dir, "formula")
    os.makedirs(formula_dir, exist_ok=True)

    document = Document(docx_path)
    image_index = 1

    # 遍历所有段落
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        for run in paragraph.runs:
            # 如果运行单元包含图片
            if contains_image(run):
                # 获取图片的二进制数据
                image_data = None
                for rel_id, rel in document.part.rels.items():
                    if "image" in rel.target_ref:
                        # 检查图片是否与当前 run 关联
                        run_xml = run._element.xml
                        if rel_id in run_xml:
                            image_data = rel.target_part.blob
                            break

                if not image_data:
                    print(f"未找到图片数据，跳过该图片")
                    continue

                # 判断图片大小（以字节为单位）
                if len(image_data) <= 1024:  # 小于1KB
                    save_dir = formula_dir
                else:
                    save_dir = output_dir

                # 获取图片文件格式
                image_format = rel.target_part.content_type.split("/")[-1]
                image_extension = "jpg" if image_format == "jpeg" else image_format

                # 保存图片到本地
                image_name = f"image_{image_index}.{image_extension}"
                image_path = os.path.join(save_dir, image_name)
                try:
                    with open(image_path, "wb") as img_file:
                        img_file.write(image_data)
                except Exception as e:
                    print(f"保存图片时发生错误: {e}")
                    continue

                # 转换 EMF/WMF 图片
                if image_extension in ["x-emf", "x-wmf"]:
                    converted_path = os.path.splitext(image_path)[0] + ".png"
                    convert_emf_or_wmf(image_data, converted_path)
                    os.remove(image_path)
                    image_path = converted_path

                # 查找图片后的图注（下一段）
                caption = None
                if paragraph_index + 1 < len(document.paragraphs):
                    next_paragraph = document.paragraphs[paragraph_index + 1]
                    if next_paragraph.text.strip():
                        caption = next_paragraph.text.strip()

                # 如果找到图注，用图注重命名图片
                try:
                    if caption:
                        safe_caption = "".join(c if c.isalnum() or c in " _-" else "_" for c in caption)
                        new_image_name = f"{safe_caption}.png"
                        new_image_path = os.path.join(save_dir, new_image_name)
                        os.rename(image_path, new_image_path)
                        print(f"图片已保存为：{new_image_path}")
                    else:
                        print(f"未找到图注，图片保存为：{image_path}")
                except Exception as e:
                    print(e)
                image_index += 1

# 使用示例
docx_path = "./files/*******.docx"   # 文档路径
output_dir = "./imgs"                # 图片输出路径
extract_images_and_captions(docx_path, output_dir)
